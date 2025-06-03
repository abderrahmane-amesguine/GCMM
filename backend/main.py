from fastapi import FastAPI, UploadFile, File, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import pandas as pd
import numpy as np
import io
import json

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Replace with your frontend URL in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory storage (replace with database in production)
gcmm_data = {
    "axes": [],
    "domains": [],
    "objectives": [],
    "globalScore": 0
}

# Color mapping for axes
axis_colors = [
    "#3366CC",  # Axis 1 - Legal (Blue)
    "#DC3912",  # Axis 2 - Technologies (Red)
    "#FF9900",  # Axis 3 - Organization (Orange)
    "#109618",  # Axis 4 - Capacity (Green)
    "#990099"   # Axis 5 - Cooperation (Purple)
]

def validate_excel_structure(df):
    """Validate the structure of the uploaded Excel file."""
    # Expected columns
    required_columns = [
        '#Axis',           # Column 0
        'Axis',       # Column 1
        '#Domain',       # Column 2
        'Domain',   # Column 3
        'Domain Description',   # Column 4
        'Obj. ID',      # Column 5
        'Objective',  # Column 6
        'Description',   # Column 7 
        'Level 1 (Ad hoc)',      # Column 8
        'Level 2 (Initiated)',      # Column 9
        'Level 3 (Defined)',      # Column 10
        'Level 4 (Managed)',      # Column 11
        'Level 5 (Optimized)',      # Column 12
        'Profil',    # Column 13
        'Target Profil',      # Column 14
        'Comment'    # Column 15
    ]
    
    # Check if we have enough columns
    if len(df.columns) < len(required_columns):
        raise HTTPException(
            status_code=400,
            detail=f"Excel file should have {len(required_columns)} columns, but found {len(df.columns)}"
        )
    
    # Check column names (case-insensitive)
    df_columns = [str(col).strip().lower() for col in df.columns]
    required_columns_lower = [col.lower() for col in required_columns]
    
    # Find missing columns
    missing_columns = []
    for i, required_col in enumerate(required_columns_lower):
        found = False
        for df_col in df_columns:
            if required_col in df_col:
                found = True
                break
        if not found:
            missing_columns.append(required_columns[i])
    if missing_columns:
        raise HTTPException(
            status_code=400,
            detail=f"Missing required columns: {', '.join(missing_columns)}"
        )
    
    return True

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Validate file extension
        if not file.filename.endswith(('.xlsx', '.xls')):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file format: {file.filename}. Please upload an Excel file (.xlsx or .xls)"
            )

        # Read Excel file content
        try:
            contents = await file.read()
        except Exception as e:
            print(f"Error reading file: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error reading file: {str(e)}. Please make sure it's a valid Excel file."
            )
        
        # Parse Excel file
        try:
            df = pd.read_excel(
                io.BytesIO(contents),
                engine='openpyxl'  # Explicitly use openpyxl for .xlsx files
            )
        except Exception as e:
            print(f"Error parsing Excel: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error parsing Excel file: {str(e)}. Please check the file format."
            )
        
        # Validate DataFrame
        if df.empty:
            raise HTTPException(
                status_code=400,
                detail="The Excel file is empty"
            )

        # Validate the structure of the Excel file
        validate_excel_structure(df)

        # Skip header row
        rows = df.iloc[0:].values.tolist()
        
        if not rows:
            raise HTTPException(
                status_code=400,
                detail="No data found in the Excel file after header row"
            )
        
        # Initialize empty lists
        axes = []
        domains = []
        objectives = []

        # Process data rows
        try:
            for row in rows:
                if not row[0]:  # Skip empty rows
                    continue
                axis_id = int(row[0]) if not pd.isna(row[0]) else None
                axis_name = row[1]
                domain_id = str(row[2]) if not pd.isna(row[2]) else None
                domain_name = row[3]
                domain_description = row[4]
                objective_id = str(row[5]) if not pd.isna(row[5]) else None
                objective_name = row[6]
                description = row[7]
                level1 = row[8]
                level2 = row[9]
                level3 = row[10]
                level4 = row[11]
                level5 = row[12]
                profile = int(row[13]) if not pd.isna(row[13]) else 0
                target_profile = int(row[14]) if not pd.isna(row[14]) else profile
                comment = row[15]

                # Process recommendations for each level
                recommendations = []
                for level in range(5):
                    recommendations.append({
                        "level": level + 1,
                        "actionable": row[16 + level] if not pd.isna(row[16 + level]) else "",
                        "strategic": row[21 + level] if not pd.isna(row[21 + level]) else ""
                    })
                
                # Add axis if it doesn't exist
                if axis_id and not any(a["id"] == axis_id for a in axes):
                    axes.append({
                        "id": axis_id,
                        "name": axis_name,
                        "score": 0,
                        "color": axis_colors[(axis_id - 1) % len(axis_colors)]
                    })
                
                # Add domain if it doesn't exist
                if domain_id and domain_name:
                    domain_key = f"{axis_id}-{domain_id}"
                    if not any(d["key"] == domain_key for d in domains):
                        domains.append({
                            "key": domain_key,
                            "id": domain_id,
                            "name": domain_name,
                            "description": domain_description,
                            "axisId": axis_id,
                            "score": 0
                        })
                
                # Add objective
                if objective_id and objective_name:
                    objectives.append({
                        "id": objective_id,
                        "name": objective_name,
                        "description": description or "",
                        "domainId": domain_id,
                        "axisId": axis_id,
                        "levels": [
                            {
                                "level": 1,
                                "description": level1,
                                "actionable": str(row[16]) if not pd.isna(row[16]) else "",
                                "strategic": str(row[21]) if not pd.isna(row[21]) else ""
                            },
                            {
                                "level": 2,
                                "description": level2,
                                "actionable": str(row[17]) if not pd.isna(row[17]) else "",
                                "strategic": str(row[22]) if not pd.isna(row[22]) else ""
                            },
                            {
                                "level": 3,
                                "description": level3,
                                "actionable": str(row[18]) if not pd.isna(row[18]) else "",
                                "strategic": str(row[23]) if not pd.isna(row[23]) else ""
                            },
                            {
                                "level": 4,
                                "description": level4,
                                "actionable": str(row[19]) if not pd.isna(row[19]) else "",
                                "strategic": str(row[24]) if not pd.isna(row[24]) else ""
                            },
                            {
                                "level": 5,
                                "description": level5,
                                "actionable": str(row[20]) if not pd.isna(row[20]) else "",
                                "strategic": str(row[25]) if not pd.isna(row[25]) else ""
                            }
                        ],
                        "profile": profile,
                        "target_profile": target_profile,
                        "comment": comment or ""
                    })
        except Exception as e:
            print(f"Error processing data: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error processing data: {str(e)}. Please check your Excel file format."
            )

        # Calculate scores and update data
        try:
            # Calculate domain scores
            for domain in domains:
                domain_objectives = [o for o in objectives 
                                if o["axisId"] == domain["axisId"] 
                                and o["domainId"] == domain["id"]]
                
                if domain_objectives:
                    domain["score"] = sum(o["profile"] for o in domain_objectives) / len(domain_objectives)
            
            # Calculate axis scores
            for axis in axes:
                axis_objectives = [o for o in objectives if o["axisId"] == axis["id"]]
                
                if axis_objectives:
                    axis["score"] = sum(o["profile"] for o in axis_objectives) / len(axis_objectives)
            
            # Calculate global score
            global_score = sum(axis["score"] for axis in axes) / len(axes) if axes else 0
            
            # Prepare radar data
            radar_data = [
                {
                    "axis": f"Axe {axis['id']}: {axis['name']}",
                    "score": axis["score"],
                    "fullMark": 5,
                    "color": axis["color"]
                }
                for axis in axes
            ]
            
            # Save processed data
            gcmm_data["axes"] = axes
            gcmm_data["domains"] = domains
            gcmm_data["objectives"] = objectives
            gcmm_data["globalScore"] = round(global_score, 1)
            gcmm_data["radarData"] = radar_data
        except Exception as e:
            print(f"Error calculating scores: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error calculating scores: {str(e)}"
            )
        
        return {
            "message": "File processed successfully",
            "filename": file.filename,
            "processedRows": len(rows),
            "axes": len(axes),
            "domains": len(domains),
            "objectives": len(objectives)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred: {str(e)}"
        )

def handle_nan_values(obj):
    if isinstance(obj, dict):
        return {key: handle_nan_values(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [handle_nan_values(item) for item in obj]
    elif isinstance(obj, (float, np.float64)) and (np.isnan(obj) or np.isinf(obj)):
        return None
    elif pd.isna(obj):
        return None
    return obj

@app.get("/api/data")
async def get_data():
    cleaned_data = handle_nan_values(gcmm_data)
    return JSONResponse(content=cleaned_data)

class ObjectiveEvaluation(BaseModel):
    objectiveId: str
    profile: int
    target_profile: int
    comment: str = ""
    recommendations: dict = {}

@app.post("/api/objectives/{objective_id}/evaluate")
async def evaluate_objective(objective_id: str, evaluation: ObjectiveEvaluation):
    try:
        # Find the objective
        objective = next((obj for obj in gcmm_data["objectives"] if obj["id"] == objective_id), None)
        if not objective:
            raise HTTPException(status_code=404, detail="Objective not found")
        # Update the objective
        objective["profile"] = evaluation.profile
        objective["target_profile"] = max(evaluation.target_profile, evaluation.profile)
        objective["comment"] = evaluation.comment

        if evaluation.recommendations:
            for level_idx, level_data in evaluation.recommendations.items():
                level_idx = int(level_idx)
                if 0 <= level_idx < len(objective["levels"]):
                    if "actionable" in level_data:
                        objective["levels"][level_idx]["actionable"] = level_data["actionable"]
                    if "strategic" in level_data:
                        objective["levels"][level_idx]["strategic"] = level_data["strategic"]

        # Update domain scores
        for domain in gcmm_data["domains"]:
            domain_objectives = [o for o in gcmm_data["objectives"] 
                            if o["axisId"] == domain["axisId"] 
                            and o["domainId"] == domain["id"]]
            
            if domain_objectives:
                domain["score"] = sum(o["profile"] for o in domain_objectives) / len(domain_objectives)
        
        # Update axis scores
        for axis in gcmm_data["axes"]:
            axis_objectives = [o for o in gcmm_data["objectives"] if o["axisId"] == axis["id"]]
            
            if axis_objectives:
                axis["score"] = sum(o["profile"] for o in axis_objectives) / len(axis_objectives)
        
        # Update global score
        global_score = sum(axis["score"] for axis in gcmm_data["axes"]) / len(gcmm_data["axes"]) if gcmm_data["axes"] else 0
        gcmm_data["globalScore"] = round(global_score, 1)

        # Update radar data
        gcmm_data["radarData"] = [
            {
                "axis": f"Axe {axis['id']}: {axis['name']}",
                "score": axis["score"],
                "fullMark": 5,
                "color": axis["color"]
            }
            for axis in gcmm_data["axes"]
        ]

        return JSONResponse(content={
            "message": "Evaluation saved successfully",
            "objective": objective
        })

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error saving evaluation: {str(e)}"
        )

@app.get("/api/export")
async def export_excel():
    """Export the current GCMM data to an Excel file."""
    try:        
        # Create a list of all records
        records = []
        for axis in gcmm_data["axes"]:
            axis_domains = [d for d in gcmm_data["domains"] if d["axisId"] == axis["id"]]
            for domain in axis_domains:
                # Get objectives for this specific domain
                domain_objectives = [o for o in gcmm_data["objectives"] 
                                  if o["domainId"] == domain["id"] and o["axisId"] == axis["id"]]
                for objective in domain_objectives:
                    record = {
                        "#Axis": axis["id"],
                        "Axis": axis["name"],
                        "#Domain": domain["id"],
                        "Domain": domain["name"],
                        "Domain Description": domain["description"],
                        "Obj. ID": objective["id"],
                        "Objective": objective["name"],
                        "Description": objective["description"],
                        "Level 1 (Ad hoc)": objective["levels"][0]["description"],
                        "Level 2 (Initiated)": objective["levels"][1]["description"],
                        "Level 3 (Defined)": objective["levels"][2]["description"],
                        "Level 4 (Managed)": objective["levels"][3]["description"],
                        "Level 5 (Optimized)": objective["levels"][4]["description"],
                        "Profil": objective["profile"],
                        "Target Profil": objective["target_profile"],
                        "Comment": objective["comment"],
                        "Actionable Recommendation for Level 1": '',
                        "Strategic Recommendation for Level 1": '',
                        "Actionable Recommendation for Level 2": '',
                        "Strategic Recommendation for Level 2": '',
                        "Actionable Recommendation for Level 3": '',
                        "Strategic Recommendation for Level 3": '',
                        "Actionable Recommendation for Level 4": '',
                        "Strategic Recommendation for Level 4": ''
                    }
                    # Only add recommendations if target_profile is greater than profile
                    if objective["target_profile"] > objective["profile"]:
                        # Add recommendations from current profile to target profile
                        for i in range(objective["profile"]-1, objective["target_profile"]):
                            if i == len(objective["levels"])-1:
                                continue
                            level = objective["levels"][i]
                            record[f"Actionable Recommendation for Level {i+1}"] = level["actionable"]
                            record[f"Strategic Recommendation for Level {i+1}"] = level["strategic"]
                    
                    records.append(record)

        # Create DataFrame and write to Excel
        df = pd.DataFrame(records)
        
        # Create an in-memory Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='GCMM')
        
        # Get the value of the BytesIO buffer
        excel_data = output.getvalue()
        
        headers = {
            'Content-Disposition': 'attachment; filename="GCMM_Export.xlsx"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }
        
        return Response(content=excel_data, headers=headers)
    
    except Exception as e:
        print(f"Error during export: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/axes/{axis_id}/export")
async def export_axis_excel(axis_id: int):
    """Export a specific axis data to an Excel file."""
    try:
        # Verify axis exists
        axis = next((a for a in gcmm_data["axes"] if a["id"] == axis_id), None)
        if not axis:
            raise HTTPException(status_code=404, detail=f"Axis {axis_id} not found")
            
        # Create a list of records for this axis only
        records = []
        axis_domains = [d for d in gcmm_data["domains"] if d["axisId"] == axis_id]
        
        for domain in axis_domains:
            # Get objectives for this specific domain
            domain_objectives = [o for o in gcmm_data["objectives"] 
                               if o["domainId"] == domain["id"] and o["axisId"] == axis_id]
            
            for objective in domain_objectives:
                record = {
                    "#Axis": axis["id"],
                    "Axis": axis["name"],
                    "#Domain": domain["id"],
                    "Domain": domain["name"],
                    "Domain Description": domain["description"],
                    "Obj. ID": objective["id"],
                    "Objective": objective["name"],
                    "Description": objective["description"],
                    "Level 1 (Ad hoc)": objective["levels"][0]["description"],
                    "Level 2 (Initiated)": objective["levels"][1]["description"],
                    "Level 3 (Defined)": objective["levels"][2]["description"],
                    "Level 4 (Managed)": objective["levels"][3]["description"],
                    "Level 5 (Optimized)": objective["levels"][4]["description"],
                    "Profil": objective["profile"],
                    "Target Profil": objective["target_profile"],
                    "Comment": objective["comment"],
                    "Actionable Recommendation for Level 1": '',
                    "Strategic Recommendation for Level 1": '',
                    "Actionable Recommendation for Level 2": '',
                    "Strategic Recommendation for Level 2": '',
                    "Actionable Recommendation for Level 3": '',
                    "Strategic Recommendation for Level 3": '',
                    "Actionable Recommendation for Level 4": '',
                    "Strategic Recommendation for Level 4": ''
                }
                
                # Add recommendations columns
                if objective["target_profile"] > objective["profile"]:
                    # Add recommendations from current profile to target profile
                    for i in range(objective["profile"]-1, objective["target_profile"]):
                        if i == len(objective["levels"])-1:
                            continue
                        level = objective["levels"][i]
                        record[f"Actionable Recommendation for Level {i+1}"] = level["actionable"]
                        record[f"Strategic Recommendation for Level {i+1}"] = level["strategic"]
                
                records.append(record)

        # Create DataFrame and write to Excel
        df = pd.DataFrame(records)
        
        # Create an in-memory Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name=f'Axis {axis_id}')
        
        # Get the value of the BytesIO buffer
        excel_data = output.getvalue()
        
        headers = {
            'Content-Disposition': f'attachment; filename="GCMM_Axis_{axis_id}_Export.xlsx"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }
        
        return Response(content=excel_data, headers=headers)
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error during axis export: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/axes/{axis_id}/report")
async def generate_axis_report(axis_id: int):
    """Generate a Word report for a specific axis."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Verify axis exists
        axis = next((a for a in gcmm_data["axes"] if a["id"] == axis_id), None)
        if not axis:
            raise HTTPException(status_code=404, detail=f"Axis {axis_id} not found")

        # Get axis data
        axis_domains = [d for d in gcmm_data["domains"] if d["axisId"] == axis_id]
        axis_objectives = [o for o in gcmm_data["objectives"] if o["axisId"] == axis_id]
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'GCMM Report - {axis["name"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add axis summary
        doc.add_heading('Axis Summary', level=1)
        doc.add_paragraph(f'Score: {axis["score"]:.1f}/5')
        
        # Add domain summaries
        doc.add_heading('Domains', level=1)
        for domain in axis_domains:
            doc.add_heading(f'{domain["name"]} (Score: {domain["score"]:.1f}/5)', level=2)
            doc.add_paragraph(domain["description"])
            
            # Add objectives for this domain
            domain_objectives = [o for o in axis_objectives if o["domainId"] == domain["id"]]
            for objective in domain_objectives:
                doc.add_heading(f'Objective {objective["id"]}: {objective["name"]}', level=3)
                doc.add_paragraph(f'Current Level: {objective["profile"]}')
                doc.add_paragraph(f'Target Level: {objective["target_profile"]}')
                doc.add_paragraph(objective["description"])
                
                # Add recommendations if target_profile > profile
                if objective["target_profile"] > objective["profile"]:
                    doc.add_heading('Recommendations', level=4)
                    for i in range(objective["profile"]-1, objective["target_profile"]):
                        if i == len(objective["levels"])-1:
                            continue
                        level = objective["levels"][i]
                        if level["actionable"] or level["strategic"]:
                            doc.add_paragraph(f'Level {i+1}:', style='Heading 5')
                            if level["actionable"]:
                                doc.add_paragraph(f'Actionable: {level["actionable"]}')
                            if level["strategic"]:
                                doc.add_paragraph(f'Strategic: {level["strategic"]}')
        
        # Save the document to a BytesIO object
        output = io.BytesIO()
        doc.save(output)
        doc_data = output.getvalue()
        
        headers = {
            'Content-Disposition': f'attachment; filename="GCMM_Axis_{axis_id}_Report.docx"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        return Response(content=doc_data, headers=headers)
    
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx library is required for generating Word reports. Please install it with 'pip install python-docx'"
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error during report generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/template")
async def download_template():
    """Download a template Excel file for GCMM data."""
    try:
        # Create template data
        template_data = [
            {
                "#Axis": "",
                "Axis": "",
                "#Domain": "",
                "Domain": "",
                "Domain Description": "",
                "Obj. ID": "",
                "Objective": "",
                "Description": "",
                "Level 1 (Ad hoc)": "",
                "Level 2 (Initiated)": "",
                "Level 3 (Defined)": "",
                "Level 4 (Managed)": "",
                "Level 5 (Optimized)": "",
                "Profil": "",
                "Target Profil": "",
                "Comment": "",
                "Actionable Recommendation for Level 1": "",
                "Strategic Recommendation for Level 1": "",
                "Actionable Recommendation for Level 2": "",
                "Strategic Recommendation for Level 2": "",
                "Actionable Recommendation for Level 3": "",
                "Strategic Recommendation for Level 3": "",
                "Actionable Recommendation for Level 4": "",
                "Strategic Recommendation for Level 4": ""
            }
        ]
        
        # Create DataFrame
        df = pd.DataFrame(template_data)
        
        # Create Excel file
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='GCMM Template')
            
            # Get the worksheet
            worksheet = writer.sheets['GCMM Template']
            
            # Adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        excel_data = output.getvalue()
        
        headers = {
            'Content-Disposition': 'attachment; filename="GCMM_Template.xlsx"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }
        
        return Response(content=excel_data, headers=headers)
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Pydantic models for request validation
class ObjectiveLevel(BaseModel):
    description: str

class Objective(BaseModel):
    id: str
    name: str
    description: str
    levels: dict[str, ObjectiveLevel]
    profile: int
    targetProfile: int
    comment: str | None = None

class Domain(BaseModel):
    id: str
    name: str
    description: str
    objectives: list[Objective] = []

class Axis(BaseModel):
    id: str
    name: str
    color: str
    description: str
    domains: list[Domain] = []

class GCMMData(BaseModel):
    axes: list[Axis]

@app.post("/api/data")
async def save_gcmm_data(data: GCMMData):
    """Save GCMM data to the backend."""
    try:
        # Update the in-memory storage
        global gcmm_data
         
        # Transform the data to match our storage format
        formatted_data = {
            "axes": [],
            "domains": [],
            "objectives": [],
            "globalScore": 0
        }
        # Process axes
        try:
            for axis in data.axes:
                axis_data = {
                    "id": str(axis.id),  # Ensure ID is string
                    "name": axis.name,
                    "color": axis.color,
                    "description": axis.description,
                    "score": 0
                }
                formatted_data["axes"].append(axis_data)

                # Process domains for this axis
                for domain in axis.domains:
                    domain_data = {
                        "id": str(domain.id),  # Ensure ID is string
                        "name": domain.name,
                        "description": domain.description,
                        "axisId": str(axis.id),  # Ensure axisId is string
                        "score": 0
                    }
                    formatted_data["domains"].append(domain_data)

                    # Process objectives for this domain
                    for objective in domain.objectives:
                        # Convert levels dict to list format expected by frontend
                        levels_list = []
                        for level_num in range(1, 6):
                            level_key = str(level_num)
                            level_data = objective.levels.get(level_key, ObjectiveLevel(description=""))
                            levels_list.append({
                                "level": level_num,
                                "description": level_data.description,
                                "actionable": "",  # Add empty actionable field
                                "strategic": ""    # Add empty strategic field
                            })

                        objective_data = {
                            "id": str(objective.id),  # Ensure ID is string
                            "name": objective.name,
                            "description": objective.description,
                            "axisId": str(axis.id),  # Ensure axisId is string
                            "domainId": str(domain.id),  # Ensure domainId is string
                            "levels": levels_list,
                            "profile": objective.profile,
                            "target_profile": objective.targetProfile,
                            "comment": objective.comment or "",
                            "score": objective.profile
                        }
                        formatted_data["objectives"].append(objective_data)
        except Exception as e:
            print(f"Error in data transformation: {str(e)}")  # Debug print
            raise

        # Calculate scores
        try:
            # Calculate domain scores
            for domain in formatted_data["domains"]:
                domain_objectives = [o for o in formatted_data["objectives"] 
                                  if o["axisId"] == domain["axisId"] 
                                  and o["domainId"] == domain["id"]]
                if domain_objectives:
                    domain["score"] = sum(o["profile"] for o in domain_objectives) / len(domain_objectives)
            
            # Calculate axis scores
            for axis in formatted_data["axes"]:
                axis_objectives = [o for o in formatted_data["objectives"] 
                                 if o["axisId"] == axis["id"]]
                if axis_objectives:
                    axis["score"] = sum(o["profile"] for o in axis_objectives) / len(axis_objectives)

            # Calculate global score
            if formatted_data["axes"]:
                formatted_data["globalScore"] = sum(axis["score"] for axis in formatted_data["axes"]) / len(formatted_data["axes"])
            
            # Add radar data
            formatted_data["radarData"] = [
                {
                    "axis": f"Axe {axis['id']}: {axis['name']}",
                    "score": axis["score"],
                    "fullMark": 5,
                    "color": axis["color"]
                }
                for axis in formatted_data["axes"]
            ]
        except Exception as e:
            print(f"Error in score calculation: {str(e)}")  # Debug print
            raise

        # Update storage
        gcmm_data.update(formatted_data)

        return JSONResponse(content={
            "message": "GCMM data saved successfully",
            "data": formatted_data
        })
    except Exception as e:
        print(f"ERROR in save_gcmm_data: {str(e)}")  # Debug print
        import traceback
        print("Traceback:", traceback.format_exc())  # Debug print full traceback
        raise HTTPException(status_code=500, detail=str(e))